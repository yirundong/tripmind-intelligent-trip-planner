# -*- coding: utf-8 -*-
"""生成 4.1 系统架构实现 Word 文档。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_run(run, en='Times New Roman', cn='宋体', size=12, bold=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, cn='黑体', size=14, bold=True)


def add_para(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run(run)


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    set_run(run, cn='黑体', bold=True, size=12)


def add_code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(10)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


add_h2('4.1 系统架构实现')

add_para(
    '本系统是一个面向智能旅行规划场景的完整 Web 应用，整体采用前后端分离的浏览器/服务器架构进行开发，'
    '包含用户认证、数据库持久化、异步任务、多智能体工作流、地图能力以及管理后台等多个组成部分，'
    '并非简单包装一次大模型调用。'
)

add_para(
    '后端基于 FastAPI 框架进行开发，使用 Pydantic 完成请求体与响应体的数据契约定义，'
    '通过 SQLAlchemy ORM 与 PostgreSQL 进行数据交互；身份认证基于 JWT 实现，密码使用加盐哈希进行存储。'
    '智能体工作流由 LangGraph 与 LangChain 协同搭建：LangGraph 负责多节点的状态管理与流程编排，'
    'LangChain 负责将工具能力以统一接口暴露给大模型；模型调用通过 OpenAI 兼容协议接入，'
    '便于在不同模型供应商之间灵活切换。地图相关能力通过高德地图开放平台的 REST 接口接入，'
    '包括兴趣点检索、路线规划、天气查询等；图片素材通过 Unsplash 服务获取。'
)

add_para(
    '前端基于 Vue 3 框架结合 TypeScript 编写，使用 Vite 作为构建工具，组件库选用 Ant Design Vue，'
    'HTTP 请求统一通过 axios 封装。地图视图直接调用高德地图的 Web JS 接口，'
    '将后端返回的结构化数据渲染为带标注与折线的可视化界面；'
    '行程结果支持通过 html2canvas 与 jsPDF 导出为图片或 PDF 文件。'
)

add_para(
    '整个项目按“前端、后端、文档、脚本”四类目录平级组织，各目录之间相互独立、互不依赖，'
    '便于团队协作与版本管理。项目整体的目录结构如图 4-1 所示。'
)

tree = '''intelligent-trip-planner/
├── backend/
│   ├── app/
│   │   ├── agents/              # LangGraph 多 Agent 工作流
│   │   ├── api/routes/          # FastAPI 路由
│   │   ├── models/              # Pydantic 请求/响应模型
│   │   ├── services/            # LLM、地图、图片服务封装
│   │   ├── tools/               # LangChain 工具
│   │   ├── auth.py              # JWT、密码哈希、当前用户
│   │   ├── database.py          # 数据库连接和 Session
│   │   └── db_models.py         # SQLAlchemy ORM 模型
│   ├── requirements.txt
│   └── .env.example
├── frontend/
│   ├── src/
│   │   ├── services/            # API 请求封装
│   │   ├── styles/              # 全局样式
│   │   ├── types/               # TypeScript 类型
│   │   └── views/               # 页面视图
│   ├── package.json
│   └── .env.example
├── docs/
│   └── ARCHITECTURE.md          # 架构图、需求分析、实现说明与论文写作底稿
├── scripts/
│   └── init_postgres.sql        # PostgreSQL 建库和授权脚本
├── start-dev.ps1                # Windows 一键启动脚本
└── README.md'''

add_code(tree)
add_caption('图 4-1  项目整体目录结构')

add_para(
    '后端项目位于 backend 目录下，所有源代码集中放置在 app 包内，按职责进一步细分为五个子包：'
)

add_para(
    '（1）agents 子包：智能体工作流的核心代码，使用 LangGraph 编排多个节点完成需求解析、景点搜索、'
    '天气查询、酒店推荐、行程生成、结构校验与结果修复等环节；'
)

add_para(
    '（2）api 子包：HTTP 接口层，main.py 完成 FastAPI 应用的初始化、跨域设置与全局异常处理；'
    'routes 子目录下按业务划分为认证、行程、收藏、地图、兴趣点和管理后台等多个路由模块；'
)

add_para(
    '（3）services 子包：业务服务层，对大模型调用、地图服务、图片素材服务等外部依赖进行统一封装，'
    '对上层屏蔽具体实现差异，内部还包含管理员账号初始化等启动期任务；'
)

add_para(
    '（4）tools 子包：将地图等能力封装为符合智能体工具调用规范的接口，供工作流节点直接使用；'
)

add_para(
    '（5）models 子包：通过 Pydantic 定义对外接口的请求体与响应体模型，'
    '与持久化模型解耦，便于在不影响数据库结构的情况下调整接口契约。'
)

add_para(
    'app 包顶层另放置四个全局基础模块：auth.py 实现 JWT 鉴权、密码哈希以及当前用户解析；'
    'database.py 管理数据库连接与会话；db_models.py 通过 SQLAlchemy 定义所有业务表；'
    'config.py 集中管理配置项。backend 目录下还包含 requirements.txt 依赖清单与 .env.example 环境变量模板。'
)

add_para(
    '前端项目位于 frontend 目录下，使用 Vue 3 + TypeScript + Vite 工具链构建，源码集中放置在 src 目录中：'
    'views 以页面为单位组织 .vue 单文件组件，涵盖登录、注册、规划、结果、我的行程、收藏、探索、'
    '路线、个人中心、控制台与管理后台等业务页面；services 中的 api.ts 基于 axios 集中封装与后端通信的请求方法；'
    'types 维护与后端字段对应的 TypeScript 类型声明；styles 存放全局样式；'
    '顶层 App.vue 与 main.ts 完成根组件挂载与路由、组件库的注册。'
    'frontend 目录下还包含 package.json 与 .env.example 环境变量模板。'
)

add_para(
    '文档目录 docs 用于集中存放系统需求分析、架构设计、设计图（drawio 格式）以及论文写作底稿等资料；'
    'scripts 目录存放数据库初始化脚本 init_postgres.sql 等运维脚本；'
    '项目根目录还提供 start-dev.ps1 一键启动脚本与 README.md 项目说明文档，便于他人快速上手与本地演示。'
    '整个项目的目录划分清晰、职责单一，每一类文件都能在固定的位置找到，便于后期的扩展与维护。'
)

out = r'c:\Users\Ariza\Desktop\毕业论文\系统架构实现.docx'
doc.save(out)
print('saved:', out)

import os
print('size:', os.path.getsize(out), 'bytes')
