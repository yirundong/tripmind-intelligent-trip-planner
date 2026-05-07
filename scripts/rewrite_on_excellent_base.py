import re
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


BASE_DOC = Path(r"c:\Users\Ariza\Desktop\毕业论文\上届参考论文.docx")
MD_SOURCE = Path(r"c:\Users\Ariza\code\intelligent-trip-planner\docs\THESIS_FINAL_BODY.md")
OUT_DOC = Path(r"c:\Users\Ariza\Desktop\毕业论文\毕业论文-终版.docx")


SECTION_MAP = [
    ("摘  要", "摘 要"),
    ("Abstract", "Abstract"),
    ("第1章  概  述", "第1章 概 述"),
    ("第2章  系统需求分析", "第2章 系统需求分析"),
    ("第3章  系统设计", "第3章 系统设计"),
    ("第4章  系统的实现与测试", "第4章 系统的实现与测试"),
    ("总  结", "总 结"),
    ("致  谢", "致 谢"),
    ("参考文献", "参考文献"),
]


def parse_md_sections(md_text: str):
    sections = {}
    current_h1 = None
    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            if current_h1:
                sections[current_h1].append(("blank", ""))
            continue
        if line.startswith("![") or line.strip() == "---":
            continue
        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                current_h1 = text
                sections.setdefault(current_h1, [])
            elif current_h1:
                sections[current_h1].append((f"h{level}", text))
            continue
        if current_h1:
            sections[current_h1].append(("p", line))
    return sections


def key_from_h1(title: str) -> str:
    t = title.strip()
    if t in ("摘  要", "摘 要"):
        return "摘 要"
    if t == "Abstract":
        return "Abstract"
    if t in ("第1章  概  述", "第1章 概 述"):
        return "第1章 概 述"
    if t in ("第2章  系统需求分析", "第2章 系统需求分析"):
        return "第2章 系统需求分析"
    if t in ("第3章  系统设计", "第3章 系统设计"):
        return "第3章 系统设计"
    if t in ("第4章  系统的实现与测试", "第4章 系统的实现与测试"):
        return "第4章 系统的实现与测试"
    if t in ("总  结", "总 结"):
        return "总 结"
    if t in ("致  谢", "致 谢"):
        return "致 谢"
    if t == "参考文献":
        return "参考文献"
    return t


def _is_section_break_paragraph(node) -> bool:
    if node.tag != qn("w:p"):
        return False
    p_pr = node.find(qn("w:pPr"))
    if p_pr is None:
        return False
    sect_pr = p_pr.find(qn("w:sectPr"))
    return sect_pr is not None


def _clear_runs_keep_pPr(p_elem):
    """Strip a paragraph of all child elements except w:pPr (which carries
    the section-break info). Effectively turns the paragraph into an empty
    one while preserving section formatting."""
    for child in list(p_elem):
        if child.tag == qn("w:pPr"):
            continue
        p_elem.remove(child)


def remove_between(current_p, next_p):
    """Remove all paragraphs/tables strictly between current_p and next_p.
    Section-break-bearing paragraphs are preserved structurally but their
    text content is cleared. Returns the LAST surviving section-break
    paragraph element if any, which should be used as the insertion anchor
    so that newly inserted content stays in the current section."""
    nodes = []
    node = current_p._element.getnext()
    while node is not None and (next_p is None or node is not next_p._element):
        nodes.append(node)
        node = node.getnext()
    last_section_anchor = None
    for n in nodes:
        if _is_section_break_paragraph(n):
            _clear_runs_keep_pPr(n)
            last_section_anchor = n
        else:
            n.getparent().remove(n)
    return last_section_anchor


def insert_blocks_before(anchor_p, blocks):
    for kind, text in blocks:
        if kind == "blank":
            anchor_p.insert_paragraph_before("", style="Normal")
            continue
        if kind == "h2":
            anchor_p.insert_paragraph_before(text, style="Heading 2")
            continue
        if kind == "h3":
            anchor_p.insert_paragraph_before(text, style="Heading 3")
            continue
        if kind == "p":
            style = "Normal"
            if text.startswith("关键词："):
                style = "Plain Text"
            anchor_p.insert_paragraph_before(text, style=style)


def clear_cover_before_first_h1(first_h1):
    """Remove everything before the first H1 except section-break paragraphs.
    Section-break paragraphs are kept structurally but their runs are cleared
    so any pre-existing cover-section page setup is preserved."""
    body = first_h1._element.getparent()
    to_remove = []
    to_clear = []
    for child in list(body):
        if child is first_h1._element:
            break
        if _is_section_break_paragraph(child):
            to_clear.append(child)
            continue
        to_remove.append(child)
    for node in to_remove:
        node.getparent().remove(node)
    for node in to_clear:
        _clear_runs_keep_pPr(node)


def prepend_blank_cover_page(first_h1):
    """Insert a single empty paragraph carrying a page-break run before the
    first H1. This produces a blank page in front of 摘要 that serves as a
    cover-page placeholder which the user can replace manually with their
    actual cover later."""
    from docx.oxml import OxmlElement
    blank = first_h1.insert_paragraph_before("", style="Normal")
    run_elem = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_elem.append(br)
    blank._element.append(run_elem)


def find_section_boundary(cur_p, next_p):
    """Walk siblings strictly after cur_p. Return the first node that acts
    as a boundary, plus the list of intermediate paragraphs/tables to be
    deleted. Boundary candidates (in priority order):
      1. next_p's element (the next H1 paragraph)
      2. a paragraph carrying a section break (w:sectPr inside w:pPr)
      3. any non-paragraph/non-table sibling, e.g. the body-level w:sectPr.
         These are body-level boundaries and must NOT be deleted."""
    intermediate = []
    node = cur_p._element.getnext()
    while node is not None:
        if next_p is not None and node is next_p._element:
            return node, intermediate
        if node.tag not in (qn("w:p"), qn("w:tbl")):
            return node, intermediate
        if _is_section_break_paragraph(node):
            return node, intermediate
        intermediate.append(node)
        node = node.getnext()
    return None, intermediate


def _insert_blocks_before_elem(boundary_elem, blocks, parent_wrap):
    """Insert blocks just before an arbitrary XML element by creating a
    temporary anchor paragraph (wrapped with the same parent as existing
    body paragraphs) and removing it after insertion."""
    from docx.oxml import OxmlElement
    anchor_elem = OxmlElement("w:p")
    boundary_elem.addprevious(anchor_elem)
    anchor_para = Paragraph(anchor_elem, parent_wrap)
    insert_blocks_before(anchor_para, blocks)
    anchor_elem.getparent().remove(anchor_elem)


def main():
    doc = Document(str(BASE_DOC))
    md_text = MD_SOURCE.read_text(encoding="utf-8")
    md_sections = parse_md_sections(md_text)

    h1_paras = [p for p in doc.paragraphs if p.style and p.style.name.lower().startswith("heading 1")]
    if len(h1_paras) < len(SECTION_MAP):
        raise RuntimeError("优秀论文底稿的一级标题数量不足，无法按节替换")

    # Cover handling: clear all paragraphs/tables before the first H1 (摘要),
    # then insert a blank page placeholder before 摘要 so the user can fill in
    # an actual cover later.
    clear_cover_before_first_h1(h1_paras[0])
    prepend_blank_cover_page(h1_paras[0])

    for i, (display_title, key_title) in enumerate(SECTION_MAP):
        cur = h1_paras[i]
        nxt = h1_paras[i + 1] if i + 1 < len(h1_paras) else None

        cur.text = display_title

        blocks = md_sections.get(key_title, [])
        if not blocks:
            blocks = md_sections.get(display_title, [])

        boundary, intermediate = find_section_boundary(cur, nxt)

        for n in intermediate:
            n.getparent().remove(n)

        if boundary is None:
            anchor = doc.add_paragraph("", style="Normal")
            insert_blocks_before(anchor, blocks)
            anchor._element.getparent().remove(anchor._element)
        elif nxt is not None and boundary is nxt._element:
            insert_blocks_before(nxt, blocks)
        elif boundary.tag == qn("w:p") and _is_section_break_paragraph(boundary):
            _clear_runs_keep_pPr(boundary)
            anchor_para = Paragraph(boundary, cur._parent)
            insert_blocks_before(anchor_para, blocks)
        else:
            _insert_blocks_before_elem(boundary, blocks, cur._parent)

    doc.save(str(OUT_DOC))
    print(f"generated: {OUT_DOC}")


if __name__ == "__main__":
    main()
