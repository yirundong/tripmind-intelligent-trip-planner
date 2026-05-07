import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BASE_DOC = Path(r"c:\Users\Ariza\Desktop\毕业论文\上届参考论文.docx")
MD_SOURCE = Path(r"c:\Users\Ariza\code\intelligent-trip-planner\docs\THESIS_DRAFT_INTEGRATED.md")
OUT_DOC = Path(r"c:\Users\Ariza\Desktop\毕业论文\毕业论文-整合稿-按优秀论文底稿-v3.docx")


H1_ORDER = [
    "摘 要",
    "Abstract",
    "目  录",
    "第1章  概  述",
    "第2章  系统需求分析",
    "第3章  系统设计",
    "第4章  系统的实现与测试",
    "总  结",
    "致 谢",
    "参考文献",
]


def normalize_h1(text: str) -> str:
    text = text.strip()
    mapping = {
        "摘 要": "摘 要",
        "摘  要": "摘 要",
        "Abstract": "Abstract",
        "目  录": "目  录",
        "目 录": "目  录",
        "第1章 概 述": "第1章  概  述",
        "第2章 系统需求分析": "第2章  系统需求分析",
        "第3章 系统设计": "第3章  系统设计",
        "第4章 系统的实现与测试": "第4章  系统的实现与测试",
        "总 结": "总  结",
        "总  结": "总  结",
        "致 谢": "致 谢",
        "参考文献": "参考文献",
    }
    return mapping.get(text, text)


def parse_markdown(md_text: str):
    sections = {}
    current_h1 = None

    for raw in md_text.splitlines():
        line = raw.rstrip()
        if not line:
            if current_h1:
                sections[current_h1].append(("blank", ""))
            continue

        # skip images and separator
        if line.startswith("![") or line.strip() == "---":
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            if level == 1:
                h1 = normalize_h1(text)
                current_h1 = h1
                sections.setdefault(current_h1, [])
            elif current_h1:
                sections[current_h1].append((f"h{level}", text))
            continue

        if current_h1:
            sections[current_h1].append(("p", line))

    return sections


def clear_body_keep_sectpr(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_pgnum_type(section, fmt=None, start=None):
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:pgNumType"):
            sect_pr.remove(child)
    if fmt is not None or start is not None:
        pg = OxmlElement("w:pgNumType")
        if fmt is not None:
            pg.set(qn("w:fmt"), fmt)
        if start is not None:
            pg.set(qn("w:start"), str(start))
        sect_pr.append(pg)


def clear_runs(paragraph):
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_toc_field(paragraph):
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "（目录，按F9更新）"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(fld_char1)
    run.append(instr)
    run.append(fld_char2)
    run.append(hint)
    run.append(fld_char3)


def append_section_content(doc: Document, h1: str, blocks):
    # heading 1
    h1_show = "摘  要" if h1 == "摘 要" else h1
    p = doc.add_paragraph(h1_show, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if h1 == "目  录":
        p2 = doc.add_paragraph(style="Normal")
        add_toc_field(p2)
        return

    for kind, text in blocks:
        if kind == "blank":
            doc.add_paragraph("", style="Normal")
            continue
        if kind == "h2":
            doc.add_paragraph(text, style="Heading 2")
            continue
        if kind == "h3":
            doc.add_paragraph(text, style="Heading 3")
            continue
        if kind == "p":
            doc.add_paragraph(text, style="Normal")


def apply_section_page_setup(doc: Document):
    header_titles = [""] + H1_ORDER
    # doc.sections includes blank cover + 10 content sections
    for i, sec in enumerate(doc.sections):
        sec.top_margin = Cm(3.3)
        sec.bottom_margin = Cm(3.0)
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.5)
        sec.header_distance = Cm(2.5)
        sec.footer_distance = Cm(2.0)
        sec.different_first_page_header_footer = False
        sec.header.is_linked_to_previous = False
        sec.footer.is_linked_to_previous = False

        hp = sec.header.paragraphs[0] if sec.header.paragraphs else sec.header.add_paragraph()
        clear_runs(hp)
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i > 0 and i < len(header_titles):
            run = hp.add_run("摘  要" if header_titles[i] == "摘 要" else header_titles[i])
            run.font.size = Pt(9)
            run.font.name = "宋体"
            r_pr = run._element.get_or_add_rPr()
            r_fonts = r_pr.find(qn("w:rFonts"))
            if r_fonts is None:
                r_fonts = OxmlElement("w:rFonts")
                r_pr.append(r_fonts)
            r_fonts.set(qn("w:eastAsia"), "宋体")

        fp = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
        clear_runs(fp)
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i > 0:
            add_page_field(fp)

    # page numbers
    # 0 blank cover no page
    set_pgnum_type(doc.sections[0], None, None)
    # 1 摘要 I
    set_pgnum_type(doc.sections[1], "upperRoman", 1)
    # 2 abstract II
    set_pgnum_type(doc.sections[2], "upperRoman", None)
    # 3 toc roman continue
    set_pgnum_type(doc.sections[3], "upperRoman", None)
    # 4 chapter1 start 1
    set_pgnum_type(doc.sections[4], "decimal", 1)
    for idx in range(5, len(doc.sections)):
        set_pgnum_type(doc.sections[idx], None, None)

    try:
        doc.settings.odd_and_even_pages_header_footer = False
    except Exception:
        pass


def main():
    base = Document(str(BASE_DOC))
    md = MD_SOURCE.read_text(encoding="utf-8")
    parsed = parse_markdown(md)

    clear_body_keep_sectpr(base)

    # blank cover page
    base.add_paragraph("", style="Normal")

    for i, h1 in enumerate(H1_ORDER):
        base.add_section(WD_SECTION_START.NEW_PAGE)
        blocks = parsed.get(h1, [])
        append_section_content(base, h1, blocks)

    apply_section_page_setup(base)
    base.save(str(OUT_DOC))
    print(f"generated: {OUT_DOC}")


if __name__ == "__main__":
    main()

