# -*- coding: utf-8 -*-
"""生成 4.3 测试章节的两个表格 Word 文档（表 4-6 测试环境 + 表 4-7 功能测试用例及结果）。"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn

doc = Document()

styles = doc.styles
normal = styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(10.5)
normal._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def set_run(run, en='Times New Roman', cn='宋体', size=10.5, bold=False):
    run.font.name = en
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), cn)


def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run(run, cn='黑体', bold=True, size=10.5)


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)

    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run(run, cn='宋体', bold=True, size=10.5)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_i, row in enumerate(rows):
        cells = t.rows[r_i + 1].cells
        for c_i, val in enumerate(row):
            cells[c_i].text = ''
            p = cells[c_i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if c_i >= 2 else WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            set_run(run, cn='宋体', size=10.5)
            cells[c_i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    doc.add_paragraph('')


add_caption('表 4-6  测试环境')
add_table(
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 10 / 11'],
        ['后端运行环境', 'Python 3.10+、FastAPI、Uvicorn'],
        ['前端运行环境', 'Node.js 18+、Vue 3、Vite'],
        ['数据库', 'PostgreSQL 14+'],
        ['外部服务', 'OpenAI 兼容 LLM 接口、高德地图 API、Unsplash API'],
        ['启动方式', 'start-dev.ps1，前端 http://localhost:5173，后端 http://localhost:8000'],
    ],
    col_widths=[3.5, 11.5],
)

add_caption('表 4-7  功能测试用例及结果')
add_table(
    ['编号', '测试项', '操作步骤', '预期结果', '实际结果'],
    [
        ['TC-01', '用户注册（正向）', '输入合法邮箱、用户名和密码并提交', '注册成功，自动登录并进入用户主页', '与预期一致'],
        ['TC-02', '用户登录（正向）', '输入正确的邮箱或用户名以及密码', '登录成功，进入工作台', '与预期一致'],
        ['TC-03', '登录失败（异常）', '输入错误密码', '提示"账号或密码错误"，留在登录页', '与预期一致'],
        ['TC-04', '越权访问后台（异常）', '普通用户访问 /admin', '跳转至 403 无权限页', '与预期一致'],
        ['TC-05', '创建规划任务（正向）', '填写城市、日期、偏好后提交规划请求', '返回任务 ID，前端显示进度条与阶段标识', '与预期一致'],
        ['TC-06', '行程生成成功（正向）', '等待轮询完成', '跳转结果页，展示每日行程、预算、天气与地图', '与预期一致'],
        ['TC-07', '行程生成失败（异常）', '在 LLM 服务异常的情况下提交规划', '任务状态置为 failed，前端显示失败提示并提供重试入口', '与预期一致'],
        ['TC-08', '编辑行程（正向）', '修改景点描述、删除一个景点并保存', '页面同步更新，数据库记录被对应修改', '与预期一致'],
        ['TC-09', '导出行程（正向）', '在结果页点击导出 PNG 或 PDF', '浏览器下载对应文件，内容与页面一致', '与预期一致'],
        ['TC-10', '查看历史行程（正向）', '进入"我的行程"页面', '显示历史行程列表，并支持按城市与状态筛选', '与预期一致'],
        ['TC-11', '删除历史行程（正向）', '选择一条记录点击删除并确认', '列表中该项移除，数据库标记为已删除', '与预期一致'],
        ['TC-12', '添加收藏地点（正向）', '手动填写地点信息并保存', '收藏列表新增该地点', '与预期一致'],
        ['TC-13', 'POI 探索并收藏（正向）', '在探索页搜索 POI 并将结果加入收藏', '地图正确显示标记，收藏写入数据库', '与预期一致'],
        ['TC-14', '路线规划（正向）', '输入起点、终点和交通方式', '返回距离与耗时，地图绘制路线折线', '与预期一致'],
        ['TC-15', '个人偏好保存（正向）', '修改默认城市、交通、住宿、偏好并保存', '保存成功，规划页默认值被同步刷新', '与预期一致'],
        ['TC-16', '必填项校验（边界）', '规划页未填城市直接提交', '前端给出"请填写城市"提示，请求未发起', '与预期一致'],
        ['TC-17', '日期合法性（边界）', '结束日期早于开始日期后提交', '给出"结束日期不能早于开始日期"提示', '与预期一致'],
        ['TC-18', 'Token 缺失（异常）', '清空 localStorage 后访问受保护页面', '自动跳转登录页', '与预期一致'],
    ],
    col_widths=[1.5, 2.6, 4.2, 4.5, 2.0],
)

out = r'c:\Users\Ariza\Desktop\毕业论文\测试章节图表.docx'
doc.save(out)
print('saved:', out)

import os
print('size:', os.path.getsize(out), 'bytes')
