from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def set_east_asia_font(run, font_name: str) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), font_name)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* MERGEFORMAT"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(text)
    run._r.append(fld_char3)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)


def add_toc_field(paragraph) -> None:
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    hint = OxmlElement("w:t")
    hint.text = "（目录域，按F9更新）"
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()._r
    run.append(fld_char1)
    run.append(instr)
    run.append(fld_char2)
    run.append(hint)
    run.append(fld_char3)


def set_pgnum_type(section, fmt=None, start=None) -> None:
    sect_pr = section._sectPr
    for child in list(sect_pr):
        if child.tag == qn("w:pgNumType"):
            sect_pr.remove(child)
    if fmt is not None or start is not None:
        pg = OxmlElement("w:pgNumType")
        if fmt is not None:
            pg.set(qn("w:fmt"), fmt)
        if start is not None:
            pg.set(qn("w:start"), str(start))
        sect_pr.append(pg)


def main() -> None:
    path = r"c:\Users\Ariza\Desktop\毕业论文\毕业论文-整合稿-自动排版-v2.docx"
    doc = Document(path)

    heading1_positions = [
        i
        for i, p in enumerate(doc.paragraphs)
        if p.style and p.style.name.lower().startswith("heading 1")
    ]
    first_h1 = heading1_positions[0] if heading1_positions else 0

    # cover formatting
    for paragraph in doc.paragraphs[:first_h1]:
        text = (paragraph.text or "").strip()
        if not text:
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        if "毕业论文" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_east_asia_font(run, "隶书")
                run.bold = True
                run.font.size = Pt(42)
        elif "基于AI-Agent" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(13)
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.bold = True
                run.font.size = Pt(22)
        elif "Design and Implementation" in text:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(22)
        elif text.startswith("学院："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_east_asia_font(run, "宋体")
                run.font.size = Pt(14)
        elif text.startswith("编号"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in paragraph.runs:
                set_east_asia_font(run, "宋体")
                run.font.size = Pt(12)
        elif text.startswith("日期："):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_east_asia_font(run, "宋体")
                run.font.size = Pt(16)

    # headings and normal body
    for paragraph in doc.paragraphs:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading 1"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(13)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.bold = True
                run.font.size = Pt(18)
        elif style_name.lower().startswith("heading 2"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(13)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.bold = True
                run.font.size = Pt(16)
        elif style_name.lower().startswith("heading 3"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(13)
            paragraph.paragraph_format.space_after = Pt(13)
            paragraph.paragraph_format.line_spacing = Pt(20)
            for run in paragraph.runs:
                set_east_asia_font(run, "黑体")
                run.bold = True
                run.font.size = Pt(15)
        else:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            paragraph.paragraph_format.first_line_indent = Cm(0.74)
            for run in paragraph.runs:
                set_east_asia_font(run, "宋体")
                run.font.size = Pt(12)

    # English abstract body formatting
    heading1_positions = [
        i
        for i, p in enumerate(doc.paragraphs)
        if p.style and p.style.name.lower().startswith("heading 1")
    ]
    if len(heading1_positions) >= 3:
        abs_start = heading1_positions[1] + 1
        toc_start = heading1_positions[2]
        for paragraph in doc.paragraphs[abs_start:toc_start]:
            if paragraph.style and paragraph.style.name.lower().startswith("heading"):
                continue
            paragraph.paragraph_format.first_line_indent = None
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

    # TOC field
    for paragraph in doc.paragraphs:
        if "目录自动生成" in (paragraph.text or ""):
            clear_paragraph(paragraph)
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = Pt(20)
            add_toc_field(paragraph)
            break

    # section settings
    header_titles = [
        "",
        "摘  要",
        "Abstract",
        "目  录",
        "第1章  概  述",
        "第2章  系统需求分析",
        "第3章  系统设计",
        "第4章  系统的实现与测试",
        "总  结",
        "致  谢",
        "参考文献",
    ]
    for i, section in enumerate(doc.sections):
        section.top_margin = Cm(3.3)
        section.bottom_margin = Cm(3.0)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(2.5)
        section.footer_distance = Cm(2.0)
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_p = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        clear_paragraph(header_p)
        header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i > 0:
            run = header_p.add_run(header_titles[i] if i < len(header_titles) else "")
            set_east_asia_font(run, "宋体")
            run.font.size = Pt(9)

        footer_p = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        clear_paragraph(footer_p)
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i > 0:
            add_page_field(footer_p)

    # page numbering
    if len(doc.sections) > 0:
        set_pgnum_type(doc.sections[0], None, None)
    if len(doc.sections) > 1:
        set_pgnum_type(doc.sections[1], "upperRoman", 1)
    if len(doc.sections) > 2:
        set_pgnum_type(doc.sections[2], "upperRoman", None)
    if len(doc.sections) > 3:
        set_pgnum_type(doc.sections[3], "upperRoman", None)
    if len(doc.sections) > 4:
        set_pgnum_type(doc.sections[4], "decimal", 1)
    for i in range(5, len(doc.sections)):
        set_pgnum_type(doc.sections[i], None, None)

    try:
        doc.settings.odd_and_even_pages_header_footer = False
    except Exception:
        pass

    doc.save(path)
    print(path, len(doc.sections))


if __name__ == "__main__":
    main()

